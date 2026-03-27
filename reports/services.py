from pathlib import Path
from tempfile import NamedTemporaryFile

from docx import Document

from .models import FamilyEducation, PoorStudent, ReportPeriod, TeacherReport
from .period_utils import get_previous_approved_report


def create_or_get_report(*, teacher, school_class, period):
    report, _ = TeacherReport.objects.get_or_create(
        teacher=teacher,
        school_class=school_class,
        period=period,
        defaults={"status": "draft", "total_students_end": 0, "has_movement": False},
    )
    return report


def sync_family_education_count(report):
    family_education, _ = FamilyEducation.objects.get_or_create(report=report)
    family_education.count = family_education.students.count()
    family_education.save(update_fields=["count", "has_family_education"])
    return family_education


def mark_recurring_poor_students(report):
    previous_report = get_previous_approved_report(report.teacher, report.school_class, report.period)
    if not previous_report or not hasattr(previous_report, "academic_performance"):
        return 0

    previous_pairs = {
        (student.full_name.strip().lower(), student.subject_code or student.subject.strip().lower())
        for student in previous_report.academic_performance.poor_students.all()
    }

    updated = 0
    if hasattr(report, "academic_performance"):
        for student in report.academic_performance.poor_students.all():
            key = (student.full_name.strip().lower(), student.subject_code or student.subject.strip().lower())
            is_recurring = key in previous_pairs
            if student.is_recurring != is_recurring:
                student.is_recurring = is_recurring
                student.save(update_fields=["is_recurring"])
                updated += 1
    return updated


def export_report_to_docx(report):
    document = Document()
    document.add_heading("Отчет классного руководителя", level=1)
    document.add_paragraph(f"Класс: {report.school_class.name}")
    document.add_paragraph(f"Период: {report.period.name}")
    document.add_paragraph(f"Классный руководитель: {report.teacher.full_name}")
    document.add_paragraph(f"Статус: {report.get_status_display()}")

    document.add_heading("1. Контингент", level=2)
    document.add_paragraph(f"Количество учеников на конец периода: {report.total_students_end}")
    document.add_paragraph(f"Очных учеников: {report.in_person_students_count}")

    if hasattr(report, "family_education"):
        document.add_paragraph(f"На семейном обучении: {report.family_education.count}")

    if report.movements.exists():
        document.add_heading("1.1 Движение учеников", level=2)
        table = document.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "Тип"
        hdr[1].text = "ФИО"
        hdr[2].text = "Откуда / куда"
        hdr[3].text = "Приказ"
        hdr[4].text = "Дата"
        for movement in report.movements.all():
            row = table.add_row().cells
            row[0].text = movement.get_movement_type_display()
            row[1].text = movement.student_name
            row[2].text = movement.target_school or movement.source_school or ""
            row[3].text = movement.order_number or ""
            row[4].text = movement.order_date.strftime("%d.%m.%Y") if movement.order_date else ""

    if hasattr(report, "academic_performance"):
        perf = report.academic_performance
        document.add_heading("2. Успеваемость", level=2)
        document.add_paragraph(f"Отличники: {perf.excellent_count}")
        document.add_paragraph(f"Ударники: {perf.good_count}")
        document.add_paragraph(f"С одной \"4\": {perf.one_four_count}")
        document.add_paragraph(f"С одной \"3\": {perf.one_three_count}")
        document.add_paragraph(f"Двоечники: {perf.poor_count}")
        document.add_paragraph(f"Не аттестованы: {perf.not_attested_count}")
        document.add_paragraph(f"Качество знаний: {report.quality_percentage}%")
        document.add_paragraph(f"Успеваемость: {report.success_percentage}%")

    tmp = NamedTemporaryFile(delete=False, suffix=".docx")
    document.save(tmp.name)
    return Path(tmp.name)


def period_display(period: ReportPeriod):
    return f"{period.name} ({period.academic_year})"
